

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Brand Colours ────────────────────────────────────────────────────────────
C_BLACK      = RGBColor(0x03, 0x05, 0x07)   # void background → dark text
C_ACID       = RGBColor(0x00, 0xC9, 0xA2)   # #00C9A2  teal accent
C_VIOLET     = RGBColor(0x7B, 0x54, 0xD4)   # #7B54D4  purple accent
C_AMBER      = RGBColor(0xD4, 0x8F, 0x0F)   # #D48F0F  amber
C_CRIMSON    = RGBColor(0xC0, 0x3A, 0x3A)   # #C03A3A  red
C_SAPPHIRE   = RGBColor(0x3A, 0x74, 0xD4)   # #3A74D4  blue
C_DARK_BG    = RGBColor(0x0D, 0x17, 0x26)   # section header fill
C_MID_GREY   = RGBColor(0x3A, 0x4A, 0x5C)   # body text
C_LIGHT_GREY = RGBColor(0xF2, 0xF5, 0xF8)   # table alt row
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_BORDER     = RGBColor(0xCC, 0xD6, 0xE0)   # table border

# ─── Page geometry (US Letter, 1" margins) ────────────────────────────────────
PAGE_W        = 12240   # DXA
PAGE_H        = 15840
MARGIN        = 1080    # 0.75 inch
CONTENT_W     = PAGE_W - 2 * MARGIN   # 10080 DXA


# ═══════════════════════════════════════════════════════════════════════════════
#  LOW-LEVEL HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    """Fill a table cell background with a solid colour (hex without #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd element first to avoid duplicates
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def _set_cell_border(cell, **edges):
    """
    Set borders on a cell.
    edges: top / bottom / left / right = dict(val, sz, color)
    """
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in edges.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   attrs.get("val",   "single"))
        el.set(qn("w:sz"),    str(attrs.get("sz", 4)))
        el.set(qn("w:color"), attrs.get("color", "auto"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_margins(cell, top=80, bottom=80, left=140, right=140):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing tcMar to avoid duplicates
    for existing in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(existing)
    mar  = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _set_cell_width(cell, width_dxa: int):
    """Safely set cell width, removing any existing w:tcW first."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:before"), str(before))
    spc.set(qn("w:after"),  str(after))
    if line:
        spc.set(qn("w:line"),     str(line))
        spc.set(qn("w:lineRule"), "auto")
    pPr.append(spc)


def _add_run(para, text, bold=False, italic=False, color=None,
             size=None, font="Calibri", underline=False):
    run        = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font
    if size:  run.font.size  = Pt(size)
    if color: run.font.color.rgb = color
    return run


def _horizontal_rule(doc, color="00C9A2"):
    """Thin coloured paragraph border used as a visual divider."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pb   = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pb.append(bot)
    pPr.append(pb)
    _para_spacing(p, before=0, after=60)
    return p


def _section_pill(doc, label: str, color_hex: str, text_hex: str = "FFFFFF"):
    """Dark label paragraph that acts like a section header pill."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex.upper())
    pPr.append(shd)
    _para_spacing(p, before=160, after=60)
    run = p.add_run(f"  {label.upper()}  ")
    run.bold           = True
    run.font.name      = "Calibri"
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(
        int(text_hex[0:2], 16), int(text_hex[2:4], 16), int(text_hex[4:6], 16)
    )
    return p


# ═══════════════════════════════════════════════════════════════════════════════
#  SCORE TABLE  (5 dimensions + total)
# ═══════════════════════════════════════════════════════════════════════════════

SCORE_DIMS = [
    ("Schema Markup",   "schema",         20, "00C9A2"),
    ("Entity Clarity",  "entity",         15, "7B54D4"),
    ("Content Depth",   "content",        25, "D48F0F"),
    ("Trust Signals",   "trust",          20, "3A74D4"),
    ("Extractability",  "extractability", 20, "C03A3A"),
]

def _build_score_table(doc, context: dict, score_data: dict):
    ss  = context.get("section_scores", {})
    av  = context.get("ai_visibility_summary", {})
    pen = context.get("penalties", {})

    # col widths:  Dimension | Score | Max | Bar | Band
    COL_W = [3200, 900, 900, 3480, 1600]   # sum = 10080

    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"

    # ── Header row ──
    hdr_texts = ["Dimension", "Score", "Max", "Progress", "Status"]
    for i, cell in enumerate(tbl.rows[0].cells):
        _set_cell_bg(cell, "0D1726")
        _cell_margins(cell)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(p, before=60, after=60)
        _add_run(p, hdr_texts[i], bold=True, color=C_WHITE, size=9)

    # ── Data rows ──
    for idx, (label, key, max_val, hex_c) in enumerate(SCORE_DIMS):
        row   = tbl.add_row()
        score = ss.get(key, 0) or 0
        pct   = round((score / max_val) * 100)
        color = RGBColor(int(hex_c[0:2],16), int(hex_c[2:4],16), int(hex_c[4:6],16))
        fill  = "F2F5F8" if idx % 2 == 0 else "FFFFFF"

        # Dimension name
        c0 = row.cells[0]
        _set_cell_bg(c0, fill)
        _cell_margins(c0)
        p0 = c0.paragraphs[0]
        _para_spacing(p0, before=60, after=60)
        _add_run(p0, label, bold=True, color=C_MID_GREY, size=10)

        # Score
        c1 = row.cells[1]
        _set_cell_bg(c1, fill)
        _cell_margins(c1)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(p1, before=60, after=60)
        _add_run(p1, str(score), bold=True, color=color, size=11)

        # Max
        c2 = row.cells[2]
        _set_cell_bg(c2, fill)
        _cell_margins(c2)
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(p2, before=60, after=60)
        _add_run(p2, str(max_val), color=C_MID_GREY, size=9)

        # Progress bar (text-based blocks)
        c3 = row.cells[3]
        _set_cell_bg(c3, fill)
        _cell_margins(c3)
        p3 = c3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_spacing(p3, before=70, after=70)
        filled  = round(pct / 5)   # out of 20 blocks
        empty   = 20 - filled
        _add_run(p3, "█" * filled, color=color,      size=8, font="Courier New")
        _add_run(p3, "░" * empty,  color=C_BORDER,   size=8, font="Courier New")
        _add_run(p3, f"  {pct}%",  color=C_MID_GREY, size=8)

        # Status badge text
        c4 = row.cells[4]
        _set_cell_bg(c4, fill)
        _cell_margins(c4)
        p4 = c4.paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(p4, before=60, after=60)
        if   pct >= 80: badge, bcol = "Excellent", RGBColor(0x00,0x9E,0x77)
        elif pct >= 60: badge, bcol = "Good",      RGBColor(0x3A,0x74,0xD4)
        elif pct >= 40: badge, bcol = "Fair",      RGBColor(0xD4,0x8F,0x0F)
        else:           badge, bcol = "Poor",      RGBColor(0xC0,0x3A,0x3A)
        _add_run(p4, badge, bold=True, color=bcol, size=9)

    # ── Penalty row (if any) ──
    if pen:
        pen_total = sum(pen.values())
        row = tbl.add_row()
        _set_cell_bg(row.cells[0], "FFF3F3")
        _cell_margins(row.cells[0])
        p = row.cells[0].paragraphs[0]
        _para_spacing(p, before=60, after=60)
        _add_run(p, "Penalties", bold=True, color=C_CRIMSON, size=10)
        for i in range(1, 5):
            _set_cell_bg(row.cells[i], "FFF3F3")
            _cell_margins(row.cells[i])
        p1 = row.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p1, str(pen_total), bold=True, color=C_CRIMSON, size=11)
        p3 = row.cells[3].paragraphs[0]
        pen_labels = ", ".join(k.replace("_"," ").title() for k in pen.keys())
        _add_run(p3, pen_labels, italic=True, color=C_CRIMSON, size=9)

    # ── Total row ──
    final = av.get("final_score", 0) or 0
    maxp  = av.get("max_possible", 100) or 100
    pct_t = round((final / maxp) * 100)
    row   = tbl.add_row()
    _set_cell_bg(row.cells[0], "0D1726")
    _cell_margins(row.cells[0])
    pt = row.cells[0].paragraphs[0]
    _para_spacing(pt, before=80, after=80)
    _add_run(pt, "TOTAL AI READINESS", bold=True, color=C_WHITE, size=10)
    for i in range(1, 5):
        _set_cell_bg(row.cells[i], "0D1726")
        _cell_margins(row.cells[i])
    p1 = row.cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p1, str(final), bold=True, color=C_ACID, size=13)
    p2 = row.cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p2, str(maxp),  color=C_WHITE, size=9)
    p3 = row.cells[3].paragraphs[0]
    filled = round(pct_t / 5)
    _add_run(p3, "█" * filled,      color=C_ACID,   size=8, font="Courier New")
    _add_run(p3, "░" * (20-filled), color=C_BORDER, size=8, font="Courier New")
    _add_run(p3, f"  {pct_t}%", bold=True, color=C_WHITE, size=9)
    p4 = row.cells[4].paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    band = av.get("readiness_band", "—")
    # short version for the cell
    short = band.split("—")[0].strip() if "—" in band else band
    _add_run(p4, short, bold=True, color=C_ACID, size=9)

    # Fix column widths on all rows
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            _set_cell_width(cell, COL_W[i])

    return tbl


# ═══════════════════════════════════════════════════════════════════════════════
#  MARKDOWN → DOCX PARSER
#  Handles the LLM output: ##headings, **bold**, `code`, ```blocks, tables, bullets
# ═══════════════════════════════════════════════════════════════════════════════

def _parse_inline(para, text: str, base_size=10.5, base_color=None):
    """Parse **bold**, *italic*, `code` inline markers into runs."""
    bc = base_color or C_MID_GREY
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts   = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            _add_run(para, part[2:-2], bold=True, color=bc, size=base_size)
        elif part.startswith("*") and part.endswith("*"):
            _add_run(para, part[1:-1], italic=True, color=bc, size=base_size)
        elif part.startswith("`") and part.endswith("`"):
            _add_run(para, part[1:-1], color=C_VIOLET, size=base_size-0.5,
                     font="Courier New")
        else:
            _add_run(para, part, color=bc, size=base_size)


def _render_code_block(doc, code: str, lang: str = ""):
    """Render a fenced code block as a light-grey shaded paragraph block."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F3F6")
    pPr.append(shd)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),  "360")
    ind.set(qn("w:right"), "360")
    pPr.append(ind)
    _para_spacing(p, before=80, after=80)
    if lang:
        _add_run(p, f"{lang}\n", bold=True, color=C_VIOLET, size=8, font="Courier New")
    _add_run(p, code.strip(), color=C_MID_GREY, size=8.5, font="Courier New")


def _render_md_table(doc, header_row: list, data_rows: list):
    """Render a markdown pipe-table as a proper docx table."""
    if not header_row:
        return
    n_cols  = len(header_row)
    col_w   = CONTENT_W // n_cols

    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    # Header
    for i, cell in enumerate(tbl.rows[0].cells):
        _set_cell_bg(cell, "0D1726")
        _cell_margins(cell)
        p = cell.paragraphs[0]
        _para_spacing(p, before=60, after=60)
        _add_run(p, header_row[i].strip(), bold=True, color=C_WHITE, size=9)

    for ri, row_data in enumerate(data_rows):
        row  = tbl.add_row()
        fill = "F2F5F8" if ri % 2 == 0 else "FFFFFF"
        for i, cell in enumerate(row.cells):
            _set_cell_bg(cell, fill)
            _cell_margins(cell)
            text = row_data[i].strip() if i < len(row_data) else ""
            p    = cell.paragraphs[0]
            _para_spacing(p, before=50, after=50)
            # Colour-code effort / score-gain cells
            low  = text.strip() == "Low"
            mid  = text.strip() == "Medium"
            gain = re.match(r'^\+\d', text.strip())
            if   low:  _add_run(p, text, bold=True, color=RGBColor(0x00,0x9E,0x77), size=9)
            elif mid:  _add_run(p, text, bold=True, color=C_AMBER, size=9)
            elif gain: _add_run(p, text, bold=True, color=C_ACID,  size=9)
            else:      _parse_inline(p, text, base_size=9)

    # Fix column widths
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            _set_cell_width(cell, col_w)

    doc.add_paragraph()   # spacing after table


def _render_markdown(doc, md_text: str):
    """
    Convert LLM markdown output to rich docx content.
    Handles: ## headings, ### headings, bullet lists, numbered lists,
             **bold**, `inline code`, ```fenced blocks```, pipe tables,
             and plain paragraphs.
    """
    if not md_text:
        return

    lines         = md_text.split("\n")
    i             = 0
    in_code_block = False
    code_lines    = []
    code_lang     = ""
    table_lines   = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = []
        for tl in table_lines:
            cells = [c.strip() for c in tl.strip("|").split("|")]
            rows.append(cells)
        # first row = header, second row = separator (skip), rest = data
        header = rows[0] if rows else []
        data   = [r for r in rows[2:] if r] if len(rows) > 2 else []
        _render_md_table(doc, header, data)
        table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Fenced code block ──
        if stripped.startswith("```"):
            if not in_code_block:
                flush_table()
                in_code_block = True
                code_lang  = stripped[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                _render_code_block(doc, "\n".join(code_lines), code_lang)
                code_lines = []
                code_lang  = ""
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Pipe table ──
        if stripped.startswith("|"):
            table_lines.append(stripped)
            i += 1
            continue
        else:
            flush_table()

        # ── Empty line ──
        if not stripped:
            i += 1
            continue

        # ── H2 heading ──
        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=200, after=60)
            pPr = p._p.get_or_add_pPr()
            pb  = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "4")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "00C9A2")
            pb.append(bot); pPr.append(pb)
            _add_run(p, text, bold=True, color=C_DARK_BG, size=12.5)
            i += 1
            continue

        # ── H3 heading ──
        if stripped.startswith("### "):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=140, after=40)
            _add_run(p, "▌ ", color=C_VIOLET, size=10, font="Courier New")
            _add_run(p, text, bold=True, color=C_DARK_BG, size=11)
            i += 1
            continue

        # ── H4 heading (####) ──
        if stripped.startswith("#### "):
            text = stripped[5:].strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=100, after=30)
            _add_run(p, text, bold=True, italic=True, color=C_MID_GREY, size=10.5)
            i += 1
            continue

        # ── Numbered list ──
        m_num = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            _para_spacing(p, before=30, after=30)
            _parse_inline(p, m_num.group(2), base_size=10)
            i += 1
            continue

        # ── Bullet list ──
        if stripped.startswith(("- ", "* ", "• ")):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            _para_spacing(p, before=30, after=30)
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"),    "360")
            ind.set(qn("w:hanging"), "200")
            pPr.append(ind)
            _add_run(p, "●  ", color=C_ACID, size=8)
            _parse_inline(p, text, base_size=10)
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r'^[-*_]{3,}$', stripped):
            _horizontal_rule(doc)
            i += 1
            continue

        # ── Plain paragraph ──
        p = doc.add_paragraph()
        _para_spacing(p, before=40, after=40, line=276)
        _parse_inline(p, stripped, base_size=10.5)
        i += 1

    flush_table()


# ═══════════════════════════════════════════════════════════════════════════════
#  QUICK-WIN / STRUCTURAL TABLE  (from prioritized_plan tables)
#  The markdown renderer handles these inline, but this is kept for direct use.
# ═══════════════════════════════════════════════════════════════════════════════

def _build_action_table(doc, rows_data: list, headers: list, col_weights: list = None):
    """Build a styled action-plan table directly from parsed row data."""
    n = len(headers)
    if col_weights is None:
        col_weights = [1] * n
    total_w  = sum(col_weights)
    col_w    = [int(CONTENT_W * w / total_w) for w in col_weights]

    tbl = doc.add_table(rows=1, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    for i, cell in enumerate(tbl.rows[0].cells):
        _set_cell_bg(cell, "0D1726")
        _cell_margins(cell)
        p = cell.paragraphs[0]
        _para_spacing(p, before=60, after=60)
        _add_run(p, headers[i], bold=True, color=C_WHITE, size=9)

    for ri, row_vals in enumerate(rows_data):
        row  = tbl.add_row()
        fill = "F2F5F8" if ri % 2 == 0 else "FFFFFF"
        for i, cell in enumerate(row.cells):
            _set_cell_bg(cell, fill)
            _cell_margins(cell)
            text = str(row_vals[i]) if i < len(row_vals) else ""
            p    = cell.paragraphs[0]
            _para_spacing(p, before=50, after=50)
            _parse_inline(p, text, base_size=9)

    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            _set_cell_width(cell, col_w[i])


# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

def _build_cover(doc, context: dict, score_data: dict):
    av = context.get("ai_visibility_summary", {})
    ps = context.get("product_summary", {})
    pi = context.get("page_identity", {})

    # Top colour bar
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "0D1726")
    pPr.append(shd)
    _para_spacing(p, before=0, after=0)
    _add_run(p, "  GENY  ·  GEO Intelligence Engine", bold=True,
             color=C_ACID, size=11)

    doc.add_paragraph()

    # Product name
    p_name = doc.add_paragraph()
    _para_spacing(p_name, before=80, after=40)
    prod_name = ps.get("name") or pi.get("title") or "Product Page"
    _add_run(p_name, prod_name, bold=True, color=C_DARK_BG, size=22)

    # Brand + price line
    p_sub = doc.add_paragraph()
    _para_spacing(p_sub, before=0, after=60)
    brand    = ps.get("brand", "")
    currency = ps.get("currency", "")
    price    = ps.get("price", "")
    curr_sym = {"USD": "$", "GBP": "£", "EUR": "€", "INR": "₹"}.get(currency, currency)
    parts = [brand, f"{curr_sym}{price}" if price else ""]
    sub_text = "  ·  ".join(x for x in parts if x)
    _add_run(p_sub, sub_text, color=C_MID_GREY, size=12)

    _horizontal_rule(doc, "00C9A2")

    # Score summary row
    final   = av.get("final_score", 0) or 0
    maxp    = av.get("max_possible", 100) or 100
    pct     = av.get("ai_readiness_pct", 0) or 0
    band    = av.get("readiness_band", "—")
    url     = pi.get("url", "")

    p_score = doc.add_paragraph()
    _para_spacing(p_score, before=60, after=60)
    _add_run(p_score, f"{final}/{maxp}  ", bold=True, color=C_ACID, size=28)
    _add_run(p_score, f"({pct}%)  ", bold=True, color=C_MID_GREY, size=14)
    _add_run(p_score, band, italic=True, color=C_VIOLET, size=12)

    if url:
        p_url = doc.add_paragraph()
        _para_spacing(p_url, before=20, after=80)
        _add_run(p_url, url, color=C_SAPPHIRE, size=9, underline=True)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  WEAK AREAS SUMMARY BOX
# ═══════════════════════════════════════════════════════════════════════════════

def _build_weak_areas(doc, context: dict):
    wa = context.get("weak_areas", {})
    if not wa:
        return
    _section_pill(doc, "⚠  Signals Scoring Zero — Must Fix", "C03A3A")

    for section, fields in wa.items():
        p = doc.add_paragraph()
        _para_spacing(p, before=40, after=20)
        _add_run(p, f"{section.title()}:  ", bold=True, color=C_CRIMSON, size=10)
        _add_run(p, "  |  ".join(f.replace("_"," ").title() for f in fields),
                 color=C_MID_GREY, size=9.5)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  PENALTIES BOX
# ═══════════════════════════════════════════════════════════════════════════════

def _build_penalties(doc, context: dict):
    pen = context.get("penalties", {})
    if not pen:
        return
    _section_pill(doc, "Penalties Applied", "7B54D4")
    for k, v in pen.items():
        p = doc.add_paragraph()
        _para_spacing(p, before=30, after=30)
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        pPr.append(ind)
        _add_run(p, k.replace("_"," ").title() + ":  ", bold=True,
                 color=C_VIOLET, size=10)
        _add_run(p, str(v), bold=True, color=C_CRIMSON, size=10)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  SCORE PROJECTION (from prioritized_plan text)
# ═══════════════════════════════════════════════════════════════════════════════

def _build_score_projection(doc, plan_text: str):
    """
    Extracts the Score Projection block from the prioritized_plan
    and renders it as a clean summary table.
    """
    pattern = re.compile(
        r'Score Projection\s*\n([\s\S]+?)(?=\n###|\n##|\Z)', re.IGNORECASE
    )
    m = pattern.search(plan_text)
    if not m:
        return

    block = m.group(1).strip()
    lines = [l.strip() for l in block.split("\n") if l.strip()]

    _section_pill(doc, "Score Projection", "00C9A2")

    tbl   = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"
    cols = [4000, 6080]

    for i, cell in enumerate(tbl.rows[0].cells):
        _set_cell_bg(cell, "0D1726")
        _cell_margins(cell)
        p = cell.paragraphs[0]
        _para_spacing(p, before=60, after=60)
        _add_run(p, ["Milestone", "Projected Score"][i],
                 bold=True, color=C_WHITE, size=9)

    for ri, line in enumerate(lines):
        # parse "Label: X / Y" or "Label: X"
        m2   = re.match(r'^(.+?):\s*(.+)$', line)
        if not m2:
            continue
        label, value = m2.group(1).strip(), m2.group(2).strip()
        row  = tbl.add_row()
        fill = "F2F5F8" if ri % 2 == 0 else "FFFFFF"
        for i, cell in enumerate(row.cells):
            _set_cell_bg(cell, fill)
            _cell_margins(cell)

        p0 = row.cells[0].paragraphs[0]
        _para_spacing(p0, before=60, after=60)
        _add_run(p0, label, bold=True, color=C_MID_GREY, size=10)

        p1 = row.cells[1].paragraphs[0]
        _para_spacing(p1, before=60, after=60)
        is_band = "band" in label.lower()
        color   = C_VIOLET if is_band else C_ACID
        _add_run(p1, value, bold=True, color=color, size=11)

    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            _set_cell_width(cell, cols[i])

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
#  HEADER / FOOTER
# ═══════════════════════════════════════════════════════════════════════════════

def _set_header_footer(section, product_name: str):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    hdr = section.header
    hdr.is_linked_to_previous = False
    hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()

    # left: brand name
    run_l = hp.add_run("GENY GEO Report  ·  ")
    run_l.bold           = True
    run_l.font.size      = Pt(8)
    run_l.font.color.rgb = C_ACID

    # right-aligned product name via tab stop
    from docx.oxml import OxmlElement
    pPr  = hp._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(CONTENT_W))
    tabs.append(tab); pPr.append(tabs)
    run_r = hp.add_run(f"\t{product_name[:60]}")
    run_r.font.size      = Pt(8)
    run_r.font.color.rgb = C_MID_GREY

    # header bottom border
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "00C9A2")
    pBdr.append(bot); pPr.append(pBdr)

    # Footer: page number
    ftr = section.footer
    ftr.is_linked_to_previous = False
    fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = fp.add_run("Page ")
    run_f.font.size      = Pt(8)
    run_f.font.color.rgb = C_MID_GREY
    # page number field
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
    fldChar2  = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run_pn = fp.add_run()
    run_pn._r.append(fldChar1); run_pn._r.append(instrText); run_pn._r.append(fldChar2)
    run_pn.font.size      = Pt(8)
    run_pn.font.color.rgb = C_MID_GREY
    run_end = fp.add_run("  ·  GENY GEO Intelligence")
    run_end.font.size      = Pt(8)
    run_end.font.color.rgb = C_MID_GREY


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

def build_report_docx(context: dict, score_data: dict, geo_result: dict) -> Document:
    """
    Build and return a polished python-docx Document object.

    Parameters
    ----------
    context    : dict — output of LLMContextBuilder.build_context()
    score_data : dict — output of AIScoringEngine.compute_score()
    geo_result : dict — the dict returned by the /geo_recommendation API endpoint
                  expected keys: executive_summary, technical_analysis,
                                 content_analysis, prioritized_plan
    """
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────────
    section       = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    ps           = context.get("product_summary", {})
    product_name = ps.get("name") or context.get("page_identity", {}).get("title") or "Product"

    _set_header_footer(section, product_name)

    # ── Default paragraph style ──────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ══════════════════════════════════════════════════════════════════════════
    #  1. COVER
    # ══════════════════════════════════════════════════════════════════════════
    _build_cover(doc, context, score_data)

    # ══════════════════════════════════════════════════════════════════════════
    #  2. AI READINESS SCORE TABLE
    # ══════════════════════════════════════════════════════════════════════════
    _section_pill(doc, "AI Readiness Score — Dimension Breakdown", "0D1726")
    _build_score_table(doc, context, score_data)
    doc.add_paragraph()

    # Penalties + weak areas
    _build_penalties(doc, context)
    _build_weak_areas(doc, context)

    # ══════════════════════════════════════════════════════════════════════════
    #  3. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    exec_summary = (
        geo_result.get("executive_summary") or
        geo_result.get("final_report") or   # backwards compat
        ""
    )
    if exec_summary:
        doc.add_paragraph()   # spacer
        _section_pill(doc, "Executive Summary", "00C9A2")
        _render_markdown(doc, exec_summary)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    #  4. TECHNICAL AUDIT
    # ══════════════════════════════════════════════════════════════════════════
    technical = geo_result.get("technical_analysis", "")
    if technical:
        _section_pill(doc, "Technical Audit Findings", "C03A3A")
        _render_markdown(doc, technical)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    #  5. CONTENT STRATEGY
    # ══════════════════════════════════════════════════════════════════════════
    content = geo_result.get("content_analysis", "")
    if content:
        _section_pill(doc, "Content Strategy Gaps", "D48F0F")
        _render_markdown(doc, content)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    #  6. PRIORITY ROADMAP  (tables rendered inline by _render_markdown)
    # ══════════════════════════════════════════════════════════════════════════
    plan = geo_result.get("prioritized_plan", "")
    if plan:
        _section_pill(doc, "Priority Implementation Roadmap", "7B54D4")
        _render_markdown(doc, plan)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    #  7. SCORE PROJECTION  (extracted from plan, rendered as clean table)
    # ══════════════════════════════════════════════════════════════════════════
    if plan:
        _build_score_projection(doc, plan)

    return doc